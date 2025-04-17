import streamlit as st
import PyPDF2
import io
import re
from wordcloud import WordCloud
import matplotlib.pyplot as plt
import pandas as pd
from sklearn.feature_extraction.text import CountVectorizer
from fpdf import FPDF
from docx import Document

# --- Page Config ---
st.set_page_config(page_title="ATS-Free Resume Scanner", layout="wide")
st.title("📄 ATS-Free Resume Scanner")

# -------------------- Utility Functions --------------------

def extract_text_from_pdf(uploaded_file):
    reader = PyPDF2.PdfReader(uploaded_file)
    text = ""
    for page in reader.pages:
        text += page.extract_text() or ""
    return text


def clean_text(text):
    # Remove special characters and normalize whitespace
    text = re.sub(r'[^a-zA-Z0-9\s]', '', text)
    text = re.sub(r'\s+', ' ', text)
    return text.strip().lower()


def extract_sections(text):
    sections = {"Experience": "", "Education": "", "Skills": "", "Projects": ""}
    current = None
    for line in text.split("\n"):
        l = line.strip()
        low = l.lower()
        if "experience" in low:
            current = "Experience"
        elif "education" in low:
            current = "Education"
        elif "skill" in low:
            current = "Skills"
        elif "project" in low:
            current = "Projects"
        elif current:
            sections[current] += l + "\n"
    return sections


def keyword_match(resume_text, job_desc):
    resume_words = set(clean_text(resume_text).split())
    job_words = set(clean_text(job_desc).split())
    common = resume_words & job_words
    missing = job_words - resume_words
    match_percent = round(len(common) / len(job_words) * 100, 2) if job_words else 0
    return match_percent, common, missing, job_words


def rule_based_suggestions_v3(resume_text, job_keywords):
    tips = []
    low = resume_text.lower()
    # Generic tips
    tips.append(("Include contact information at the top (email, phone, LinkedIn)", False))
    tips.append(("Add a concise summary or objective statement at the beginning", False))
    tips.append(("Use clear section headings (Experience, Education, Skills, Projects)", False))
    tips.append(("Use bullet points to list achievements and responsibilities", False))
    # Existing rule-based checks
    if len(resume_text.split("\n")) < 10:
        tips.append(("Ensure the resume is at least 1 page long", False))
    if "experience" not in low:
        tips.append(("Add an 'Experience' section", False))
    if "education" not in low:
        tips.append(("Include an 'Education' section", False))
    if "skill" not in low:
        tips.append(("Add a 'Skills' section listing relevant tools and technologies", False))
    if "project" not in low:
        tips.append(("Add a 'Projects' section to showcase your work", False))
    # Achievement quantification
    if not re.search(r"\d+%", resume_text) and not re.search(r"\$\d+", resume_text):
        tips.append(("Quantify your achievements with numbers or metrics", False))
    # Action verbs
    if not any(w in low for w in ["led", "built", "created", "improved"]):
        tips.append(("Use strong action verbs like 'Led', 'Created', 'Improved'", False))
    # Job keywords
    if job_keywords:
        missing_jk = job_keywords - set(low.split())
        if missing_jk:
            tips.append((f"Include these job keywords: {', '.join(list(missing_jk)[:5])}...", False))
    # Final success if no issues
    if all(done for _, done in tips):
        tips = [("🎉 Your resume is well-formatted and comprehensive!", True)]
    return tips

# -------------------- Export Report Functions --------------------

def export_report_to_pdf(score, matched, missing, section_matches, suggestions):
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=14)
    pdf.cell(0, 10, "ATS Resume Scanner Report", ln=True)
    pdf.ln(5)
    pdf.set_font("Arial", size=12)
    pdf.cell(0, 8, f"Match Score: {score}%", ln=True)
    pdf.ln(3)
    pdf.multi_cell(0, 6, "Matched Keywords: " + ", ".join(list(matched)))
    pdf.ln(1)
    pdf.multi_cell(0, 6, "Missing Keywords: " + ", ".join(list(missing)))
    pdf.ln(5)
    pdf.cell(0, 8, "Section-wise Matches:", ln=True)
    for sec, words in section_matches.items():
        pdf.multi_cell(0, 6, f" - {sec}: {len(words)} matched ({', '.join(words)})")
    pdf.ln(5)
    pdf.cell(0, 8, "Resume Improvement Tips:", ln=True)
    for tip, _ in suggestions:
        pdf.multi_cell(0, 6, f" - {tip}")
    return pdf.output(dest='S').encode('latin-1')


def export_report_to_docx(score, matched, missing, section_matches, suggestions):
    doc = Document()
    doc.add_heading('ATS Resume Scanner Report', 0)
    doc.add_paragraph(f"Match Score: {score}%")
    doc.add_paragraph("Matched Keywords: " + ", ".join(list(matched)))
    doc.add_paragraph("Missing Keywords: " + ", ".join(list(missing)))
    doc.add_heading('Section-wise Matches', level=1)
    for sec, words in section_matches.items():
        doc.add_paragraph(f"{sec}: {len(words)} matched ({', '.join(words)})")
    doc.add_heading('Resume Improvement Tips', level=1)
    for tip, _ in suggestions:
        doc.add_paragraph(f"- {tip}")
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()

# -------------------- Main App --------------------

uploaded_file = st.file_uploader("📤 Upload Your Resume (PDF)", type=["pdf"])
job_description = st.text_area("📝 Paste the Job Description (optional)")

if uploaded_file:
    resume_bytes = uploaded_file.read()
    resume_text = extract_text_from_pdf(io.BytesIO(resume_bytes))
    clean_resume = clean_text(resume_text)

    match_percent, matched, missing, job_keywords = keyword_match(resume_text, job_description)
    sections = extract_sections(resume_text)
    section_matches = {}
    for sec, cont in sections.items():
        section_matches[sec] = set(clean_text(cont).split()) & job_keywords

    suggestions = rule_based_suggestions_v3(resume_text, job_keywords)

    tab1, tab2, tab3 = st.tabs(["📊 Overview", "📈 Sections", "📥 Report & Tips"])

    with tab1:
        col1, col2 = st.columns([1, 1])
        with col1:
            st.subheader("✅ Match Score")
            st.markdown(f"<h2 style='color:green'>{match_percent}%</h2>", unsafe_allow_html=True)
            st.subheader("❌ Missing Keywords")
            if missing:
                st.warning(", ".join(list(missing)[:20]) + ("..." if len(missing) > 20 else ""))
            else:
                st.success("All job keywords found!")
        with col2:
            st.subheader("☁️ Resume Word Cloud")
            wc = WordCloud(width=800, height=300, background_color='white', colormap='viridis', max_words=100).generate(clean_resume)
            fig, ax = plt.subplots(figsize=(8, 3))
            ax.imshow(wc, interpolation='bilinear'); ax.axis('off')
            st.pyplot(fig)

    with tab2:
        st.subheader("📈 Section-wise Matches")
        df_sec = pd.DataFrame([{"Section": sec, "Matches": len(words)} for sec, words in section_matches.items()])
        fig2, ax2 = plt.subplots(figsize=(6, 4))
        ax2.bar(df_sec['Section'], df_sec['Matches'], color='#4CAF50')
        ax2.set_ylabel('Matches'); ax2.set_title('Section-wise Keyword Matches')
        for i, v in enumerate(df_sec['Matches']):
            ax2.text(i, v + 0.2, str(v), ha='center')
        st.pyplot(fig2)

        st.subheader("🔢 Section Details")
        cols = st.columns(2)
        for idx, (sec, words) in enumerate(section_matches.items()):
            with cols[idx % 2]:
                st.markdown(f"**{sec}** ({len(words)} matches):")
                for w in words:
                    st.write(f"- {w}")

    with tab3:
        st.subheader("📥 Download Analysis Report")
        pdf_report = export_report_to_pdf(match_percent, matched, missing, section_matches, suggestions)
        st.download_button("Export Report as PDF", data=pdf_report, file_name="analysis_report.pdf", mime="application/pdf")
        docx_report = export_report_to_docx(match_percent, matched, missing, section_matches, suggestions)
        st.download_button("Export Report as DOCX", data=docx_report, file_name="analysis_report.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        st.subheader("🎯 Improvement Tips Checklist")
        for tip, done in suggestions:
            st.checkbox(tip, value=done, disabled=True)
